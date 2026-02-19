"""
File parsers for different document formats using decorator pattern.
Supports: CSV, PDF, DOCX, TXT, XLSX
"""

from typing import Callable
from pathlib import Path
from pypdf import PdfReader
from docx import Document
import pandas as pd
from tqdm import tqdm


# Registry for file parsers
FILE_PARSERS: dict[str, Callable] = {}

def register_parser(file_extension: str):
    """
    Decorator to register file parsers for specific extensions.
    
    Usage:
        @register_parser('.pdf')
        def pdf_parser(file_path: str) -> List[str]:
            # parser implementation
            pass
    """
    def decorator(func: Callable): 
        FILE_PARSERS[file_extension.lower()] = func
        return func
    return decorator


def get_parser(file_path: str) -> Callable:
    '''Get appropriate parser based on file extension'''
    extension = Path(file_path).suffix.lower()
    parser = FILE_PARSERS.get(extension)

    if parser is None:
        raise ValueError(f'No parser registered for file type: {extension}')
    return parser


# =================== Text Cleaning Decorators =================

def clean_whitespace(func: Callable) -> Callable:
    '''Decorator to clean extra whiespace from extracted text.'''
    def wrapper(*args, **kwargs) -> list[str]:
        texts = func(*args, **kwargs)
        return [' '.join(text.split()) for text in texts if text.strip()]
    return wrapper


def remove_short_texts(min_length: int = 20):
    '''Decorator to filter out very short text chunks'''
    def decorator(func: Callable) -> Callable:
        def wrapper(*args, **kwargs) -> list[str]:
            texts = func(*args, **kwargs)
            return [text for text in texts if len(text) >= min_length]
        return wrapper
    return decorator


def chunk_text(chunk_size: int = 500):
    '''Decorator to split long texts into smaller chunks'''
    def decorator(func: Callable) -> Callable:
        def wrapper(*args, **kwargs) -> list[str]:
            texts = func(*args, **kwargs)
            chunks = []
            for text in texts:
                # split text into chunks of specified size
                words = text.split()
                for i in range(0, len(words), chunk_size):
                    chunk = ' '.join(words[i:i + chunk_size])
                    chunks.append(chunk)
            return chunks
        return wrapper
    return decorator



# ================== File Parser Implementations ===================

@register_parser('.txt')
@clean_whitespace
@remove_short_texts(min_length=10)
def parse_text(file_path: str) -> list[str]:
    '''Parse plain text files'''
    print(f'Parsing TXT: {file_path}')

    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        content = f.read()

    # split by paragraphs (double newlines)
    paragraphs = content.split('\n\n')
    return [p.strip() for p in paragraphs if p.strip()]


@register_parser('.pdf')
@clean_whitespace
@remove_short_texts(min_length=30)
def parse_pdf(file_path: str) -> list[str]:
    '''Parse pdf files using pypdf'''
    print(f'Parsing PDF: {file_path}')

    texts = []
    try:
        reader = PdfReader(file_path)

        for page_num, page in enumerate(reader.pages):
            text = page.extract_text() or ''
            if text.strip():
                texts.append(f'Page {page_num + 1}: {text}')

        print(f'Extracted {len(texts)} pages')
    except Exception as e:
        print(f'Error parsing PDF: {e}')
    return texts


@register_parser('.docx')
@clean_whitespace
@remove_short_texts(min_length=20)
def parse_docx(file_path: str) -> list[str]:
    '''Parse docx files using python-docx'''
    print(f'Parsing DOCX: {file_path}')

    texts = []
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            if para.text.strip():
                texts.append(para.text)
        
        # also text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    texts.append(row_text)

        print(f'Extracted {len(texts)} paragraphs/rows')
    except Exception as e:
        print(f'Error parsing DOCX: {e}')
    
    return texts


@register_parser('.csv')
@clean_whitespace
@remove_short_texts(min_length=10)
def parse_csv(file_path: str) -> list[str]:
    '''Parse CSV files using pandas'''
    print(f'Parsing CSV: {file_path}')

    texts = []
    try:
        df = pd.read_csv(file_path)
        # convert each row to a text representation
        for idx, row in df.iterrows():
            row_text = ' | '.join(f'{col}: {val}' for col, val in row.items())
            texts.append(row_text)
        
        print(f"Extracted {len(texts)} rows")
    except Exception as e:
        print(f"Error parsing CSV: {e}")
    
    return texts


@register_parser('.xlsx')
@register_parser('.xls')
@clean_whitespace
@remove_short_texts(min_length=10)
def parse_excel(file_path: str) -> list[str]:
    '''Parse Excel files using pandas'''
    print(f'Parsing Excel: {file_path}')

    texts = []
    try:
        excel_file = pd.ExcelFile(file_path)

        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            # Include sheet context in each row instead of separate document
            for idx, row in df.iterrows():
                row_text = f'[Sheet: {sheet_name}] ' + ' | '.join(f'{col}: {val}' for col, val in row.items())
                texts.append(row_text)

        print(f'Extracted {len(texts)} rows from {len(excel_file.sheet_names)} sheets')
    except Exception as e:
        print(f'Error parsing Excel: {e}')
    return texts


# =================== Batch Processing ===================

def process_directory(directory_path: str, recursive: bool = False) -> dict[str, list[str]]:
    """
    Process all supported files in a directory.
    
    Args:
        directory_path: Path to directory containing files
        recursive: Whether to search subdirectories
    
    Returns:
        Dictionary mapping file paths to extracted texts
    """
    results = {}
    path = Path(directory_path)
    # Determine file pattern
    pattern = "**/*" if recursive else "*"

    for file_path in tqdm(path.glob(pattern), desc='📂 Scanning files'):
        if file_path.is_file():
            extension = file_path.suffix.lower()
            if extension in FILE_PARSERS:
                try:
                    parser = get_parser(str(file_path))
                    texts = parser(str(file_path))
                    results[str(file_path)] = texts
                except Exception as e:
                    print(f'Failed to process {file_path}: {e}')
    
    return results


def get_supported_extensions() -> list[str]:
    """Return list of supported file extensions."""
    return list(FILE_PARSERS.keys())